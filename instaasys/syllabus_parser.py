import re
import os
from docx import Document


def _clean(text: str) -> str:
    return re.sub(r'\s+', ' ', text or '').strip()


# ─── PDF support ──────────────────────────────────────────────────────────────

def _extract_pdf_text(pdf_path: str) -> str:
    """Extract all text from a PDF file using pdfplumber."""
    import pdfplumber
    lines = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                lines.append(text)
    return '\n'.join(lines)


def _extract_pdf_tables(pdf_path: str) -> list:
    """
    Extract tables from a PDF file using pdfplumber.
    Returns a list of tables, where each table is a list of rows.
    """
    import pdfplumber
    all_tables = []
    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            tables = page.extract_tables()
            for table in tables:
                if table and len(table) > 0:
                    # Clean up table data
                    cleaned_table = []
                    for row in table:
                        cleaned_row = [_clean(cell) if cell else '' for cell in row]
                        cleaned_table.append(cleaned_row)
                    all_tables.append(cleaned_table)
    return all_tables


def _parse_pdf_tables_to_weekly_plan(tables: list) -> list:
    """
    Parse PDF tables to extract weekly course plan.
    This is a best-effort extraction for PDF files.
    """
    weekly = []
    
    for table in tables:
        if not table or len(table) < 2:
            continue
        
        # Check if this looks like a course plan table
        header = ' '.join(table[0]).lower()
        if not any(keyword in header for keyword in ['week', 'time', 'topic', 'learning']):
            continue
        
        # Try to parse rows
        for row in table[1:]:
            if len(row) < 3:
                continue
            
            time_raw = _clean(row[0])
            if not time_raw or len(time_raw) < 2:
                continue
            
            # Extract week number
            week_m = re.search(r'Week\s*([\d\s\-–]+)', time_raw, re.I)
            week = _clean(week_m.group(1)).replace('–', '-') if week_m else ''
            
            # Extract label
            exam_m = re.search(r'(PRELIM|MIDTERM|PRE.?FINAL|FINAL)\s*EXAMINATION', time_raw, re.I)
            label = exam_m.group(0).upper() if exam_m else ''
            
            weekly.append({
                'week': week,
                'label': label,
                'cilos': _parse_cilos(row[1] if len(row) > 1 else ''),
                'topics': _parse_topics(row[2] if len(row) > 2 else ''),
                'performance_indicators': _parse_cilos(row[3] if len(row) > 3 else ''),
                'methodology': _parse_methodology(row[4] if len(row) > 4 else ''),
                'resources': [],
                'assessment': _parse_assessment(row[5] if len(row) > 5 else ''),
                'clo_codes': [c.upper() for c in re.findall(r'CLO\d+', row[6] if len(row) > 6 else '', re.I)],
            })
    
    return weekly


def _parse_pdf_fields(raw_text: str) -> dict:
    """Best-effort extraction of labeled fields from PDF raw text."""
    result = {
        'code': '', 'title': '', 'prerequisite': '',
        'credit_units': '', 'hours': '', 'description': '',
        'semester': '1st', 'school_year': '',
        'performance_target': '', 'gad_themes': '', 'grading_system': ''
    }

    # Course code - more flexible patterns
    patterns = [
        r'COURSE\s+CODE[:\s]+(.+?)(?:\n|COURSE\s+DESCRIPTIVE)',
        r'Course\s+Code[:\s]+(.+?)(?:\n)',
        r'CODE[:\s]+([A-Z]{2,4}\s*\d{2,4})',
    ]
    for pattern in patterns:
        m = re.search(pattern, raw_text, re.I)
        if m:
            result['code'] = _clean(m.group(1))
            break

    # Title - multiple patterns
    patterns = [
        r'COURSE\s+DESCRIPTIVE\s+TITLE[:\s]+(.+?)(?:\n|COURSE\s+PRE)',
        r'Course\s+Title[:\s]+(.+?)(?:\n)',
        r'DESCRIPTIVE\s+TITLE[:\s]+(.+?)(?:\n)',
    ]
    for pattern in patterns:
        m = re.search(pattern, raw_text, re.I)
        if m:
            result['title'] = _clean(m.group(1))
            break

    # Prerequisite - handle "None" or "N/A"
    patterns = [
        r'COURSE\s+PRE.?REQUISITE[:\s]+(.+?)(?:\n|CREDIT)',
        r'Pre.?requisite[:\s]+(.+?)(?:\n)',
        r'PRE.?REQ[:\s]+(.+?)(?:\n)',
    ]
    for pattern in patterns:
        m = re.search(pattern, raw_text, re.I)
        if m:
            prereq = _clean(m.group(1))
            result['prerequisite'] = prereq if prereq.lower() not in ['none', 'n/a', 'na'] else ''
            break

    # Credit units
    patterns = [
        r'CREDIT\s+UNITS?[:\s]+(.+?)(?:\n|NO\.?\s+OF)',
        r'Units?[:\s]+(\d+)',
        r'(\d+)\s+units?',
    ]
    for pattern in patterns:
        m = re.search(pattern, raw_text, re.I)
        if m:
            result['credit_units'] = _clean(m.group(1))
            break

    # Contact hours - handle various formats
    patterns = [
        r'NO\.?\s+OF\s+(?:CONTACT\s+)?HOURS?[:\s]+(.+?)(?:\n|PERFORMANCE)',
        r'Contact\s+Hours?[:\s]+(.+?)(?:\n)',
        r'Hours?[:\s]+(.+?)(?:\n)',
    ]
    for pattern in patterns:
        m = re.search(pattern, raw_text, re.I)
        if m:
            result['hours'] = _clean(m.group(1))
            break

    # Semester and school year
    m = re.search(r'(1st|2nd|Summer)[\s\S]{0,30}(?:Semester)?[\s\S]{0,20}(?:A\.?Y\.?|S\.?Y\.?)\s*([\d]{4}\s*[-–]\s*[\d]{4})',
                  raw_text, re.I)
    if m:
        sem_map = {'1st': '1st', '2nd': '2nd', 'summer': 'summer'}
        result['semester'] = sem_map.get(m.group(1).lower(), '1st')
        result['school_year'] = re.sub(r'\s', '', m.group(2)).replace('–', '-')

    # Description - capture multi-line, stop before references, requirements, performance target, or GAD themes
    patterns = [
        r'COURSE\s+DESCRIPTION[:\s]+(.+?)(?:COURSE\s+REFERENCES|COURSE\s+REQUIREMENTS|SUPPLEMENTAL\s+READINGS|PERFORMANCE\s+TARGET|GAD\s+THEMES|COURSE\s+LEARNING|GRADING|$)',
        r'Description[:\s]+(.+?)(?:REFERENCES|REQUIREMENTS|PERFORMANCE|LEARNING|$)',
    ]
    for pattern in patterns:
        m = re.search(pattern, raw_text, re.I | re.S)
        if m:
            desc = _clean(m.group(1))
            # Remove common artifacts
            desc = re.sub(r'\(CMO.*?\)', '', desc, flags=re.I)
            result['description'] = desc[:1000]  # Limit length
            break

    # Performance Target
    patterns = [
        r'PERFORMANCE\s+TARGET[:\s]+(.+?)(?:\n|GAD\s+THEMES|COURSE\s+DESCRIPTION|GRADING)',
        r'Performance\s+Target[:\s]+(.+?)(?:\n)',
    ]
    for pattern in patterns:
        m = re.search(pattern, raw_text, re.I | re.S)
        if m:
            result['performance_target'] = _clean(m.group(1))[:500]
            break

    # GAD Themes Integrated
    patterns = [
        r'GAD\s+THEMES?\s+INTEGRATED[:\s]+(.+?)(?:\n|COURSE\s+DESCRIPTION|GRADING|$)',
        r'GAD\s+THEMES?[:\s]+(.+?)(?:\n)',
        r'Gender.*?Themes?[:\s]+(.+?)(?:\n)',
    ]
    for pattern in patterns:
        m = re.search(pattern, raw_text, re.I | re.S)
        if m:
            result['gad_themes'] = _clean(m.group(1))[:500]
            break

    # Grading System - capture detailed breakdown, stop before policies or requirements
    patterns = [
        r'GRADING\s+SYSTEM[:\s]+(.+?)(?:COURSE\s+POLICIES|COURSE\s+REQUIREMENTS|DETAILED\s+COURSE|COURSE\s+LEARNING|$)',
        r'Grading[:\s]+(.+?)(?:POLICIES|REQUIREMENTS|COURSE\s+LEARNING|$)',
    ]
    for pattern in patterns:
        m = re.search(pattern, raw_text, re.I | re.S)
        if m:
            grading = _clean(m.group(1))
            # Clean up common artifacts
            grading = re.sub(r'Page\s+\d+', '', grading, flags=re.I)
            result['grading_system'] = grading[:1000]
            break

    return result


def _leading_spaces(line: str) -> int:
    return len(line) - len(line.lstrip(' \t'))


def _parse_topics(raw: str) -> list:
    """
    Each non-empty line is a separate topic.
    Exception: 4+ leading spaces AND previous topic ends with
    a conjunction ('and'/'or') means it's a word-wrapped continuation.
    """
    result = []
    for line in raw.split('\n'):
        stripped = _clean(line)
        if not stripped or len(stripped) < 2:
            continue
        stripped = re.sub(r'^[❖•\-\*]\s*', '', stripped)
        if not stripped:
            continue
        is_continuation = (
            _leading_spaces(line) >= 4
            and result
            and re.search(r'\b(and|or)\s*$', result[-1], re.I)
        )
        if is_continuation:
            result[-1] = result[-1] + ' ' + stripped
        else:
            result.append(stripped)
    return result


def _parse_assessment(raw: str) -> list:
    """3+ leading spaces = continuation. Hyphen at end of line = split word."""
    result = []
    for line in raw.split('\n'):
        stripped = _clean(line)
        if not stripped or len(stripped) < 3:
            continue
        if result and result[-1].endswith('-'):
            result[-1] = result[-1] + stripped
        elif _leading_spaces(line) >= 3 and result:
            result[-1] = result[-1].rstrip() + ' ' + stripped
        else:
            result.append(stripped)
    return result


def _parse_methodology(raw: str) -> list:
    """Same as assessment; skip Face-to-Face / Remote Teaching section headers."""
    result = []
    for line in raw.split('\n'):
        stripped = _clean(line)
        if not stripped or len(stripped) < 4:
            continue
        if re.match(r'Face.to.Face|Remote\s+Teaching', stripped, re.I):
            continue
        if result and result[-1].endswith('-'):
            result[-1] = result[-1] + stripped
        elif _leading_spaces(line) >= 3 and result:
            result[-1] = result[-1].rstrip() + ' ' + stripped
        else:
            result.append(stripped)
    return [r for r in result
            if not re.match(r'^(tools?|Use of offline)\s*$', r, re.I)]


def _parse_cilos(raw: str) -> list:
    """3+ leading spaces = continuation."""
    result = []
    for line in raw.split('\n'):
        stripped = _clean(line)
        if not stripped or len(stripped) < 10:
            continue
        if _leading_spaces(line) >= 3 and result:
            result[-1] = result[-1].rstrip() + ' ' + stripped
        else:
            result.append(stripped)
    return result


def _parse_header(paragraphs: list) -> dict:
    """Extract semester and school year from the header line near the top."""
    semester_map = {'1st': '1st', '2nd': '2nd', 'summer': 'summer'}
    semester, school_year = '1st', ''
    for p in paragraphs[:10]:
        t = _clean(p.text)
        m = re.search(
            r'(1st|2nd|Summer)[\s\S]{0,20}A\.?Y\.?\s*([\d]{4}\s*[-–]\s*[\d]{4})',
            t, re.IGNORECASE
        )
        if m:
            semester    = semester_map.get(m.group(1).lower(), '1st')
            school_year = re.sub(r'\s', '', m.group(2)).replace('–', '-')
            break
    return {'semester': semester, 'school_year': school_year}


def _parse_labeled_fields(paragraphs: list) -> dict:
    """Scan paragraphs for the labeled course fields."""
    LABELS = {
        'code':         re.compile(r'COURSE\s+CODE', re.I),
        'title':        re.compile(r'COURSE\s+DESCRIPTIVE\s+TITLE', re.I),
        'prerequisite': re.compile(r'COURSE\s+PRE.?REQUISITE', re.I),
        'credit_units': re.compile(r'CREDIT\s+UNITS', re.I),
        'hours':        re.compile(r'NO\.?\s+OF\s+(?:CONTACT\s+)?HOURS', re.I),
        'performance_target': re.compile(r'PERFORMANCE\s+TARGET', re.I),
        'gad_themes':   re.compile(r'GAD\s+THEMES?\s+INTEGRATED', re.I),
        'grading_system': re.compile(r'GRADING\s+SYSTEM', re.I),
    }
    result = {k: '' for k in LABELS}
    description_lines = []
    grading_lines = []
    capturing_description = False
    capturing_grading = False

    for p in paragraphs:
        t = _clean(p.text)
        if not t:
            continue
        
        matched = False
        
        # Check for labeled fields
        for key, pattern in LABELS.items():
            if pattern.match(t):
                # Extract value after colon or on same line
                if ':' in t:
                    result[key] = t.split(':', 1)[1].strip()
                else:
                    # Value might be on next line
                    result[key] = ''
                capturing_description = False
                capturing_grading = False
                matched = True
                break
        
        if matched:
            continue
        
        # Handle Course Description (multi-line)
        if re.match(r'COURSE\s+DESCRIPTION', t, re.I):
            capturing_description = True
            capturing_grading = False
            inline = t.split(':', 1)[1].strip() if ':' in t else ''
            if inline:
                description_lines.append(inline)
            continue
        
        # Handle Grading System (multi-line)
        if re.match(r'GRADING\s+SYSTEM', t, re.I):
            capturing_grading = True
            capturing_description = False
            inline = t.split(':', 1)[1].strip() if ':' in t else ''
            if inline:
                grading_lines.append(inline)
            continue
        
        # Stop capturing description when we hit these sections
        if re.match(r'COURSE\s+LEARNING\s+OUTCOMES?|COURSE\s+REFERENCES?|COURSE\s+REQUIREMENTS?|SUPPLEMENTAL\s+READINGS?|PERFORMANCE\s+TARGET|GAD\s+THEMES?', t, re.I):
            capturing_description = False
            if not re.match(r'GRADING\s+SYSTEM', t, re.I):
                capturing_grading = False
            continue
        
        # Stop capturing grading when we hit course learning plan, policies, or requirements
        if re.match(r'DETAILED\s+COURSE\s+LEARNING\s+PLAN|COURSE\s+POLICIES|COURSE\s+REQUIREMENTS', t, re.I):
            capturing_grading = False
            continue
        
        # Capture description lines
        if capturing_description and t and not re.match(r'\(CMO', t, re.I):
            # Skip page numbers and headers
            if not re.match(r'Page\s+\d+', t, re.I):
                description_lines.append(t)
        
        # Capture grading lines
        if capturing_grading and t:
            # Skip page numbers
            if not re.match(r'Page\s+\d+', t, re.I):
                grading_lines.append(t)

    # Join multi-line fields
    if description_lines:
        result['description'] = ' '.join(description_lines)
    if grading_lines:
        result['grading_system'] = ' '.join(grading_lines)
    
    # Handle prerequisite "None" cases
    if result['prerequisite'].lower() in ['none', 'n/a', 'na', 'nil']:
        result['prerequisite'] = ''
    
    return result


def _parse_clos(tables: list) -> list:
    """Find the CLO table and extract rows."""
    for table in tables:
        if not table.rows:
            continue
        if not re.search(r'course\s+learning\s+outcomes?',
                         _clean(table.rows[0].cells[0].text), re.I):
            continue
        clos = []
        for row in table.rows[1:]:
            cells = [_clean(c.text) for c in row.cells]
            if len(cells) < 3 or not cells[0]:
                continue
            raw  = cells[0]
            code = re.match(r'(CLO\d+)', raw, re.I)
            clos.append({
                'code':        code.group(1).upper() if code else '',
                'description': re.sub(r'^CLO\d+[\.\s]*', '', raw, flags=re.I).strip(),
                'ilo_codes':   [i.upper() for i in re.findall(r'ILO\d+', cells[1], re.I)],
                'plo_codes':   re.findall(r'[A-Z]{2,4}\d{2,3}', cells[2]),
            })
        return clos
    return []


def _parse_weekly_plan(tables: list) -> list:
    """
    Find the Detailed Course Learning Plan table and parse all rows.
    Extracts: TIME FRAME, TOPICS, LEARNING OUTCOMES (CILOs), 
    PERFORMANCE INDICATORS, INSTRUCTIONAL METHODOLOGY, ASSESSMENT/ACTIVITIES
    """
    for table in tables:
        if not table.rows or len(table.columns) < 5:
            continue
        
        # Check if this is the course learning plan table
        header_text = ' '.join(_clean(c.text) for c in table.rows[0].cells)
        if not re.search(r'time\s*frame|course\s*intended|detailed.*plan', header_text, re.I):
            continue

        weekly = []
        for row in table.rows[1:]:
            cells = [c.text for c in row.cells]
            
            # Skip empty rows
            if not any(_clean(c) for c in cells):
                continue
            
            time_raw = _clean(cells[0])
            if not time_raw:
                continue

            # Extract week number
            week_m = re.search(r'Week\s*([\d\s\-–]+)', time_raw, re.I)
            week   = _clean(week_m.group(1)).replace('–', '-') if week_m else ''

            # Extract exam labels
            exam_m = re.search(
                r'(PRELIM|MIDTERM|PRE.?FINAL|FINAL)\s*EXAMINATION', time_raw, re.I)
            lbl_m  = re.search(
                r'(PRELIM|MIDTERM|PRE.?FINAL|FINAL)\b', time_raw, re.I)
            label  = (
                exam_m.group(1).upper() + ' EXAMINATION' if exam_m
                else lbl_m.group(0).upper() if lbl_m
                else ''
            )

            # Parse each column
            # Column indices may vary, so we handle flexibly
            cilos_text = cells[1] if len(cells) > 1 else ''
            topics_text = cells[2] if len(cells) > 2 else ''
            perf_indicators_text = cells[3] if len(cells) > 3 else ''  # Performance Indicators
            methodology_text = cells[4] if len(cells) > 4 else ''  # Instructional Methodology
            resources_text = cells[5] if len(cells) > 5 else ''  # Resources (optional)
            assessment_text = cells[6] if len(cells) > 6 else ''  # Assessment/Activities
            clo_codes_text = cells[7] if len(cells) > 7 else ''  # CLO codes

            weekly.append({
                'week':        week,
                'label':       label,
                'cilos':       _parse_cilos(cilos_text),  # Learning Outcomes
                'topics':      _parse_topics(topics_text),
                'performance_indicators': _parse_cilos(perf_indicators_text),  # Performance Indicators
                'methodology': _parse_methodology(methodology_text),
                'resources':   _parse_methodology(resources_text),  # Resources
                'assessment':  _parse_assessment(assessment_text),
                'clo_codes':   [c.upper()
                                for c in re.findall(
                                    r'CLO\d+',
                                    clo_codes_text,
                                    re.I)],
            })
        return weekly
    return []


# ─── Public API ───────────────────────────────────────────────────────────────

def process_syllabus_file(file_path: str) -> dict:
    """
    Parse a syllabus file (.docx or .pdf) and return a structured dict.
    DOCX: full structured extraction (tables, paragraphs).
    PDF:  best-effort regex extraction from raw text + table extraction.
    
    Returns dict with keys:
    - code, title, prerequisite, credit_units, hours
    - description, semester, school_year
    - performance_target, gad_themes, grading_system
    - clos (list of CLO dicts)
    - weekly_plan (list of week dicts)
    - raw_text (full extracted text)
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        raw_text = _extract_pdf_text(file_path)
        fields = _parse_pdf_fields(raw_text)
        
        # Try to extract tables for CLOs and weekly plan
        pdf_tables = _extract_pdf_tables(file_path)
        weekly_plan = _parse_pdf_tables_to_weekly_plan(pdf_tables)
        
        return {
            **fields,
            'clos':        [],  # PDF CLO extraction is complex, left empty for now
            'weekly_plan': weekly_plan,
            'raw_text':    raw_text,
        }

    # Default: .docx
    doc = Document(file_path)
    return {
        **_parse_labeled_fields(doc.paragraphs),
        **_parse_header(doc.paragraphs),
        'clos':        _parse_clos(doc.tables),
        'weekly_plan': _parse_weekly_plan(doc.tables),
        'raw_text':    '\n'.join(
            p.text for p in doc.paragraphs if p.text.strip()
        ),
    }