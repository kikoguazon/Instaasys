# assessments/assessment_builder.py

from io import BytesIO
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Unicode circled capital letters — looks exactly like a hand-circled answer
_CIRCLED = {
    'A': 'Ⓐ', 'B': 'Ⓑ', 'C': 'Ⓒ', 'D': 'Ⓓ', 'E': 'Ⓔ',
    'F': 'Ⓕ', 'G': 'Ⓖ', 'H': 'Ⓗ', 'I': 'Ⓘ', 'J': 'Ⓙ',
}

def _circle(letter: str) -> str:
    """Return the circled Unicode version of a letter, e.g. 'A' → 'Ⓐ'."""
    return _CIRCLED.get(letter.upper().strip(), letter)


def add_border(paragraph, **kwargs):
    """Add border to paragraph."""
    p = paragraph._element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    
    for border_name in ('top', 'left', 'bottom', 'right'):
        border_el = OxmlElement(f'w:{border_name}')
        border_el.set(qn('w:val'), 'single')
        border_el.set(qn('w:sz'), '4')
        border_el.set(qn('w:space'), '1')
        border_el.set(qn('w:color'), '000000')
        pBdr.append(border_el)
    
    pPr.append(pBdr)


def build_assessment_docx(question_set, questions, course, include_answers=True):
    """
    Build a professional assessment document in DOCX format.
    
    Args:
        question_set: QuestionSet instance
        questions: List of Question instances
        course: Course instance
        include_answers: Boolean, whether to include answer key
    
    Returns:
        BytesIO object containing the DOCX file
    """
    doc = Document()
    
    # Set document margins (1 inch all around)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # ─── Header: Institution ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Republic of the Philippines\n')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    
    run = p.add_run('North Eastern Mindanao State University\n')
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run.font.bold = True
    
    run = p.add_run('Tandag City, Surigao del Sur')
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    
    doc.add_paragraph()  # Spacing
    
    # ─── Student Info Fields ──────────────────────────────────────────────
    # Name and Date on same line
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run('Name: ')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.bold = True
    
    run = p.add_run('_' * 40)
    run.font.size = Pt(11)
    
    run = p.add_run('     Date: ')
    run.font.size = Pt(11)
    run.font.bold = True
    
    run = p.add_run('_' * 20)
    run.font.size = Pt(11)
    
    # Program and Score on same line
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    
    run = p.add_run('Program: ')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.bold = True
    
    run = p.add_run('_' * 35)
    run.font.size = Pt(11)
    
    run = p.add_run('     Score: ')
    run.font.size = Pt(11)
    run.font.bold = True
    
    run = p.add_run('_' * 20)
    run.font.size = Pt(11)
    
    # ─── Directions ───────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    
    run = p.add_run('I.     Directions: ')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.bold = True
    
    # Determine question type for directions
    if questions:
        q_type = questions[0].question_type
        if q_type == 'multiple_choice':
            directions = 'Choose the BEST answer for the following questions. Circle only the letter of the correct answer.'
        elif q_type == 'true_false':
            directions = 'Write TRUE if the statement is correct, FALSE if it is incorrect.'
        elif q_type == 'identification':
            directions = 'Identify what is being asked in each item. Write your answer on the space provided.'
        elif q_type == 'essay':
            directions = 'Answer the following questions comprehensively. Write your answers in the space provided.'
        else:
            directions = 'Answer the following questions.'
    else:
        directions = 'Answer the following questions.'
    
    run = p.add_run(directions)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    
    doc.add_paragraph()  # Spacing
    
    # ─── Exam Title ───────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    
    title = question_set.title.upper()
    run = p.add_run(title)
    run.font.size = Pt(12)
    run.font.name = 'Arial'
    run.font.bold = True
    
    doc.add_paragraph()  # Spacing
    
    # ─── Questions ────────────────────────────────────────────────────────
    # Group questions by type
    mc_questions = [q for q in questions if q.question_type == 'multiple_choice']
    tf_questions = [q for q in questions if q.question_type == 'true_false']
    id_questions = [q for q in questions if q.question_type == 'identification']
    essay_questions = [q for q in questions if q.question_type == 'essay']
    oral_questions = [q for q in questions if q.question_type == 'oral']
    
    # Multiple Choice Questions (2-column layout)
    if mc_questions:
        _add_multiple_choice_section(doc, mc_questions, include_answers)
    
    # True/False Questions
    if tf_questions:
        _add_true_false_section(doc, tf_questions, include_answers)
    
    # Identification Questions
    if id_questions:
        _add_identification_section(doc, id_questions, include_answers)
    
    # Essay Questions
    if essay_questions:
        _add_essay_section(doc, essay_questions, include_answers)
    
    # Oral Questions
    if oral_questions:
        _add_oral_section(doc, oral_questions, include_answers)
    
    # ─── Answer Key (if included) ─────────────────────────────────────────
    if include_answers:
        doc.add_page_break()
        _add_answer_key(doc, questions)
    
    # Save to BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def _add_multiple_choice_section(doc, questions, include_answers):
    """Add multiple choice questions in 2-column layout."""
    # Add table with 2 columns
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.allow_autofit = False
    
    # Set column widths
    for col in table.columns:
        col.width = Inches(3.0)
    
    # Add questions in pairs
    for i in range(0, len(questions), 2):
        row = table.add_row()
        
        # Left column
        left_cell = row.cells[0]
        _format_mc_question(left_cell, questions[i], i + 1, include_answers)
        
        # Right column (if exists)
        if i + 1 < len(questions):
            right_cell = row.cells[1]
            _format_mc_question(right_cell, questions[i + 1], i + 2, include_answers)


def _format_mc_question(cell, question, number, include_answers):
    """Format a single multiple choice question in a table cell."""
    # Question number and text
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(6)
    
    run = p.add_run(f'{number}. ')
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    run.font.bold = True
    
    run = p.add_run(question.content)
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    
    # Choices
    if question.choices:
        for letter, text in sorted(question.choices.items()):
            p = cell.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(3)
            
            # Circle the correct answer if include_answers
            if include_answers and letter == question.answer_key:
                run = p.add_run(f'⊙ {letter}. ')
                run.font.bold = True
            else:
                run = p.add_run(f'{letter}. ')
            
            run.font.size = Pt(10)
            run.font.name = 'Arial'
            
            run = p.add_run(text)
            run.font.size = Pt(10)
            run.font.name = 'Arial'


def _add_true_false_section(doc, questions, include_answers):
    """Add true/false questions."""
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('II. TRUE or FALSE')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.bold = True
    
    for i, question in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        
        run = p.add_run(f'{i}. ')
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.bold = True
        
        run = p.add_run(question.content)
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        
        if include_answers:
            answer = 'TRUE' if question.answer_key == 'A' else 'FALSE'
            run = p.add_run(f'  [{answer}]')
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 128, 0)


def _add_identification_section(doc, questions, include_answers):
    """Add identification questions."""
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('III. IDENTIFICATION')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.bold = True
    
    for i, question in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        
        run = p.add_run(f'{i}. ')
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.bold = True
        
        run = p.add_run(question.content)
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        
        # Answer line
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_after = Pt(12)
        
        if include_answers:
            run = p.add_run(f'Answer: {question.answer_key}')
            run.font.size = Pt(10)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 128, 0)
        else:
            run = p.add_run('_' * 50)
            run.font.size = Pt(10)


def _add_essay_section(doc, questions, include_answers):
    """Add essay questions."""
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('IV. ESSAY')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.bold = True
    
    for i, question in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        
        run = p.add_run(f'{i}. ')
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.bold = True
        
        run = p.add_run(question.content)
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        
        # Space for answer
        for _ in range(5):
            p = doc.add_paragraph()
            run = p.add_run('_' * 80)
            run.font.size = Pt(10)
        
        if include_answers and question.expected_answer:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(f'Expected Answer: {question.expected_answer}')
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0, 128, 0)


def _add_oral_section(doc, questions, include_answers):
    """Add oral questions."""
    doc.add_paragraph()
    
    p = doc.add_paragraph()
    run = p.add_run('V. ORAL ASSESSMENT')
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    run.font.bold = True
    
    for i, question in enumerate(questions, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        
        run = p.add_run(f'{i}. ')
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run.font.bold = True
        
        run = p.add_run(question.content)
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        
        if include_answers:
            if question.answer_key:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(f'Key Points: {question.answer_key}')
                run.font.size = Pt(9)
                run.font.italic = True
                run.font.color.rgb = RGBColor(0, 128, 0)
            
            if question.rubric:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                run = p.add_run(f'Rubric: {question.rubric}')
                run.font.size = Pt(9)
                run.font.italic = True


def _add_answer_key(doc, questions):
    """Add answer key section."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('ANSWER KEY')
    run.font.size = Pt(14)
    run.font.name = 'Arial'
    run.font.bold = True
    
    doc.add_paragraph()
    
    # Group by type
    mc_questions = [(i+1, q) for i, q in enumerate(questions) if q.question_type == 'multiple_choice']
    
    if mc_questions:
        p = doc.add_paragraph()
        run = p.add_run('Multiple Choice:')
        run.font.size = Pt(11)
        run.font.bold = True
        
        # Create answer key in columns
        answers_per_row = 5
        for i in range(0, len(mc_questions), answers_per_row):
            p = doc.add_paragraph()
            row_answers = mc_questions[i:i+answers_per_row]
            answer_text = '     '.join([f'{num}. {q.answer_key}' for num, q in row_answers])
            run = p.add_run(answer_text)
            run.font.size = Pt(10)
            run.font.name = 'Courier New'
